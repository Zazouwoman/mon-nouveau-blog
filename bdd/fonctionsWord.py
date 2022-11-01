
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa

from django.conf import settings
from django.contrib.staticfiles import finders
from django.shortcuts import redirect
from .models import *
import os

from docx import Document
from docx.shared import Inches
from docx.shared import Pt,Cm
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING

def relance_word(k,facture,ingeprev,affaire,mission,payeur,envoi):

    doc = Document()
    font = doc.styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(12)
    font = doc.styles['Footer'].font
    font.name = 'Calibri'
    font.size = Pt(8)

    doc = Adresses_Word(facture,ingeprev,affaire,mission,payeur,envoi,doc)
    if k == 2:
        doc = Corps_Message2(facture,ingeprev,affaire,mission,payeur,envoi,doc)
    elif k == 3:
        doc = Corps_Message3(facture, ingeprev, affaire, mission, payeur, envoi, doc)
    elif k == 4:
        doc = Corps_Message4(facture, ingeprev, affaire, mission, payeur, envoi, doc)
    doc = Pied_Page(facture,ingeprev,affaire,mission,payeur,envoi,doc)
    return doc

def Corps_Message4(facture, ingeprev, affaire, mission, payeur, envoi, doc):
    self = facture

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(30)
    p1.paragraph_format.space_after = Pt(0)
    run = p1.add_run()
    run.add_text("Objet : Lettre recommandée de mise en demeure avec AR")
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(30)
    run = p1.add_run()
    text = "Courrier RAR n° {}".format(self.Num_RAR_Demeure)
    run.add_text(text)

    paragraphe = ['Madame, Monsieur']
    tuple = self.Numero_Facture, self.Date_Facture_aff(), self.Date_Echeance1_aff(), self.Date_Relance1_aff(), self.Num_Suivi, self.Date_Relance2_aff(), self.Num_RAR, self.Date_Relance3_aff()
    text = """Nous constatons avec regret que votre société n'a toujours pas réglé le solde de notre facture n° {} en date du {} arrivée à échéance le {} malgré nos trois relances précédentes ; une première relance par mail en date du {}, une deuxième relance par courrier suivi n°{} en date du {} puis une troisième relance par courrier RAR n° {} en date du {}.
    """.format(self.Numero_Facture, self.Date_Facture_aff(), self.Date_Echeance1_aff(), self.Date_Relance1_aff(), self.Num_Suivi, self.Date_Relance2_aff(), self.Num_RAR, self.Date_Relance3_aff())
    paragraphe.append(text)
    if self.Nb_Avoir() == 1:
        avoirlie = self.Avoirs_Lies()[0]
        montantavoir = self.Montants_Avoirs_Lies_TTC()[0]
        text = "En tenant compte de l'avoir partiel n° {} d'un montant de {} € TTC, la somme restant à régler est de {} € TTC.".format(
            avoirlie, montantavoir, self.Reste_A_Payer_TTC())
        paragraphe.append(text)
    if self.Nb_Avoir() > 1:
        avoirlie = self.Avoirs_Lies()
        montantavoir = self.Montants_Avoirs_Lies_TTC()
        text = "En tenant compte des avoirs partiels "
        for x in avoirlie:
            text += 'n°' + str(x) + ', '
        text += " de montants respectifs "
        for x in montantavoir:
            text += str(x) + ' € TTC, '
        text += "la somme restant à régler est de {} € TTC.".format(self.Reste_A_Payer_TTC())
        paragraphe.append(text)
    text = """Ainsi, par la présente, nous vous mettons en demeure de nous verser, à titre principal, la somme de {} € TTC. Cette somme sera majorée des intérêts au taux légal dus en vertu de l'aticle 1 153 du code civil. Nous vous informons que ces pénalités courent dès réception de cette lettre.""".format(self.Reste_A_Payer_TTC())
    paragraphe.append(text)
    text = """Si, dans un délai de quinze jours à compter de cette date, vous ne vous êtes toujours pas acquité de votre obligation, nous saisirons la juridiction compétente afin d'obtenie le paiement des sommes susvisées."""
    paragraphe.append(text)
    text = "Veuillez agréer, Madame, Monsieur, l'expression de nos salutations distinguées."
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(10)
        run = p1.add_run()
        run.add_text(text)

    doc = Signature(facture, ingeprev, affaire, mission, payeur, envoi, doc)

    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[6]

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(40)
    p1.paragraph_format.space_after = Pt(10)
    run = p1.add_run()
    text = "Pièces jointes : copie de la facture n° : {}".format(self.Numero_Facture)
    if self.Nb_Avoir() == 1 :
        text += " + copie de l'avoir"
    elif self.Nb_Avoir() > 1:
        text += " + copie des avoirs"
    text += " + copie de la lettre de relance n°2 + copie de la lettre de relance 3"
    run.add_text(text)

    return doc

def Corps_Message3(facture, ingeprev, affaire, mission, payeur, envoi, doc):
    self = facture

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(30)
    p1.paragraph_format.space_after = Pt(0)
    run = p1.add_run()
    run.add_text("Objet : 3ème relance suite au non-paiement d'une facture.")
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(30)
    run = p1.add_run()
    text = "Courrier RAR n° {}".format(self.Num_RAR)
    run.add_text(text)


    paragraphe = ['Madame, Monsieur']
    tuple = (self.Numero_Facture, self.Montant_Facture_TTC(), self.Num_Suivi, self.Date_Relance2_aff(), self.Date_Relance1_aff())
    text = "Nous vous avions demandé de procéder au paiement de la facture n° {} d'un montant de {} € TTC par courrier suivi n°{}, en date du {}. Pour mémoire, cette facture avait déjà dû faire l'objet d'une relance par mail en date du {}. ".format(self.Numero_Facture, self.Montant_Facture_TTC(), self.Num_Suivi, self.Date_Relance2_aff(), self.Date_Relance1_aff())
    paragraphe.append(text)
    if self.Nb_Avoir() == 1:
        avoirlie = self.Avoirs_Lies()[0]
        montantavoir = self.Montants_Avoirs_Lies_TTC()[0]
        text = "En tenant compte de l'avoir partiel n° {} d'un montant de {} € TTC, la somme restant à régler est de {} € TTC.".format(
            avoirlie, montantavoir, self.Reste_A_Payer_TTC())
        paragraphe.append(text)
    if self.Nb_Avoir() > 1:
        avoirlie = self.Avoirs_Lies()
        montantavoir = self.Montants_Avoirs_Lies_TTC()
        text = "En tenant compte des avoirs partiels "
        for x in avoirlie:
            text += 'n°' + str(x) + ', '
        text += " de montants respectifs "
        for x in montantavoir:
            text += str(x) + ' € TTC, '
        text += "la somme restant à régler est de {} € TTC.".format(self.Reste_A_Payer_TTC())
        paragraphe.append(text)
    text = "Ainsi, sauf erreur de notre part, nous sommes au regret de constater qu'aucun paiement n'a été effectué pour régler cette facture. Nous vous demandons de bien vouloir régler cette situation, et de procéder au règlement de la somme due."
    paragraphe.append(text)
    text = "Si vous avez procédé au règlement avant la reception de ce courrier, veuillez ne pas tenir compte de cette lettre."
    paragraphe.append(text)
    text = "Nous vous prions d'agréer, Madame, Monsieur, l'expression de nos salutations distinguées."
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(10)
        run = p1.add_run()
        run.add_text(text)

    doc = Signature(facture, ingeprev, affaire, mission, payeur, envoi, doc)

    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[6]

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(40)
    p1.paragraph_format.space_after = Pt(10)
    run = p1.add_run()
    text = "Pièces jointes : copie de la facture n° : {}".format(self.Numero_Facture)
    if self.Nb_Avoir() == 1 :
        text += " + copie de l'avoir"
    elif self.Nb_Avoir() > 1:
        text += " + copie des avoirs"
    text += " + copie de la lettre de relance n°2"
    run.add_text(text)

    return doc


def Corps_Message2(facture,ingeprev,affaire,mission,payeur,envoi,doc):
     self = facture

     p1 = doc.add_paragraph()
     p1.paragraph_format.space_before = Pt(30)
     p1.paragraph_format.space_after = Pt(30)
     run = p1.add_run()
     run.add_text("Objet : Lettre de relance suite au non-paiement d'une facture")

     paragraphe = ['Madame, Monsieur']
     text = "Après vérification de votre compte client et sauf erreur de notre part, nous n'avons pas perçu le règlement de la facture suivante :"
     paragraphe.append(text)
     text = "Facture n° {} d'un montant de {} € en date du {}, arrivée à échéance le {}".format(self.Numero_Facture,
                                                                                                self.Montant_Facture_TTC(),
                                                                                                self.Date_Facture_aff(),
                                                                                                self.Date_Echeance1_aff())
     paragraphe.append(text)
     if self.Nb_Avoir() == 1:
         avoirlie = self.Avoirs_Lies()[0]
         montantavoir = self.Montants_Avoirs_Lies_TTC()[0]
         text = "En tenant compte de l'avoir partiel n° {} d'un montant de {} € TTC, la somme restant à régler est de {} € TTC.".format(
             avoirlie, montantavoir, self.Reste_A_Payer_TTC())
         paragraphe.append(text)
     if self.Nb_Avoir() > 1:
         avoirlie = self.Avoirs_Lies()
         montantavoir = self.Montants_Avoirs_Lies_TTC()
         text = "En tenant compte des avoirs partiels "
         for x in avoirlie:
             text += 'n°' + str(x) + ', '
         text += " de montants respectifs "
         for x in montantavoir:
             text += str(x) + ' € TTC, '
         text += "la somme restant à régler est de {} € TTC.".format(self.Reste_A_Payer_TTC())
         paragraphe.append(text)
     text = "Pour mémoire, cette facture a déjà dû faire l'objet d'une relance par mail le {}".format(
         self.Date_Relance1_aff())
     paragraphe.append(text)
     text = "Nous vous saurions gré de bien vouloir régulariser cette situation rapidement et de nous confirmer la date du règlement."
     paragraphe.append(text)
     text = "Dans le cas où votre paiement nous parviendrait avant réception de la présente, nous vous prions de ne pas en tenir compte."
     paragraphe.append(text)
     text = "Nous restons à votre disposition pour tout complément d'information."
     paragraphe.append(text)
     for text in paragraphe:
         p1 = doc.add_paragraph()
         p1.paragraph_format.space_after = Pt(10)
         run = p1.add_run()
         run.add_text(text)

     doc = Signature(facture,ingeprev,affaire,mission,payeur,envoi,doc)

     new_section = doc.add_section(WD_SECTION.CONTINUOUS)
     new_section.start_type
     section = doc.sections[6]

     sectPr = section._sectPr
     cols = sectPr.xpath('./w:cols')[0]
     cols.set(qn('w:num'), '1')

     p1 = doc.add_paragraph()
     p1.paragraph_format.space_before = Pt(40)
     p1.paragraph_format.space_after = Pt(10)
     run = p1.add_run()
     text = "Pièces jointes : copie de la facture n° : {}".format(self.Numero_Facture)
     run.add_text(text)

     return doc


def Pied_Page(facture,ingeprev,affaire,mission,payeur,envoi,doc):
    self = facture
    # Calling the footer
    footer = doc.sections[0].footer

    # Calling the paragraph already present in
    # the footer section
    # footer_para = footer.paragraphs[0]
    para = []
    text = 'INGEPREV - {} {} {}\n'.format(ingeprev.Adresse, ingeprev.CP, ingeprev.Ville)
    para.append(text)
    text = '{} au capital de {} € - {} RCS PARIS \n '.format(ingeprev.Type_Societe, ingeprev.Capital,
                                                             ingeprev.SIRET)
    para.append(text)
    text = '{} - {} - {} - {}'.format(ingeprev.Site_Web, ingeprev.Facebook, ingeprev.Twitter, ingeprev.Linkedin)
    para.append(text)
    for text in para:
        footer_para = footer.add_paragraph()
        run = footer_para.add_run()
        run.font.size = Pt(8)
        footer_para.paragraph_format.space_after = Pt(0)
        footer_para.paragraph_format.space_before = Pt(0)
        run.add_text(text)
    # Adding text in the footer
    # font = footer_para.paragraph_format.font
    # font.size = Pt(8)
    # footer_para.text = text
    return doc


def Signature(facture,ingeprev,affaire,mission,payeur,envoi,doc):
    self = facture
    # Signature
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[4]

    # Set to 2 column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    new_section = doc.add_section(WD_SECTION.NEW_COLUMN)
    new_section.start_type

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(10)
    p1.paragraph_format.space_after = Pt(00)
    run = p1.add_run()
    run.add_text("Pour INGEPREV")
    paragraphe = ["{} {}".format(self.Prenom_Pilote, self.Nom_Pilote)]
    text = "Tél : {}".format(self.Tel_Portable_Pilote)
    paragraphe.append(text)
    text = "Mail : {}".format(self.Email_Pilote)
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    return doc

def Adresses_Word(facture,ingeprev,affaire,mission,payeur,envoi,doc):
    # Entête
    #new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    #new_section.start_type
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    # doc.sections[0].top_margin = Cm(5)

    image = ingeprev.logo
    logo_run = paragraph.add_run()
    # logo_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run.add_picture(image, width=Inches(1.5))
    header.add_paragraph('')
    # doc.add_picture(image, width=Inches(1.5))
    # last_paragraph = doc.paragraphs[-1]
    # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Second Section Entête avec dernières adresses enregistrées dans la base
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[1]

    # Set to 2 column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    new_section = doc.add_section(WD_SECTION.NEW_COLUMN)
    new_section.start_type

    if not adresses_identiques2(payeur, envoi):
        paragraphe = []
        text = '{}'.format(payeur.Denomination_Sociale)
        paragraphe.append(text)
        text = '{}'.format(payeur.Adresse)
        paragraphe.append(text)
        if payeur.Complement_Adresse != '':
            text = '{}'.format(payeur.Complement_Adresse)
            paragraphe.append(text)
        text = '{} {}'.format(payeur.CP, payeur.Ville)
        paragraphe.append(text)
        for text in paragraphe:
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            run = p1.add_run()
            run.add_text(text)

        p1 = doc.add_paragraph()
        p1.paragraph_format.space_before = Pt(20)
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text('s/c')

    paragraphe = []
    text = '{}'.format(envoi.Denomination_Sociale)
    paragraphe.append(text)
    text = '{}'.format(envoi.Adresse)
    paragraphe.append(text)
    if envoi.Complement_Adresse != '':
        text = '{}'.format(envoi.Complement_Adresse)
        paragraphe.append(text)
    text = '{} {}'.format(envoi.CP, envoi.Ville)
    paragraphe.append(text)
    if envoi.Nom_Contact != '':
        text = "A l'attention de {} {} {}".format(envoi.Civilite, envoi.Nom_Contact, envoi.Prenom_Contact)
        paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[3]

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

    paragraphe = []
    text = 'N° de dossier : {}'.format(affaire.Ref_Affaire)
    paragraphe.append(text)
    text = 'Affaire : {}'.format(mission.Nom_Mission)
    paragraphe.append(text)
    text = '{}'.format(mission.Adresse)
    paragraphe.append(text)
    if mission.Complement_Adresse != '':
        text = '{}'.format(mission.Complement_Adresse)
        paragraphe.append(text)
    text = '{} {}'.format(mission.CP, mission.Ville)
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    # p1.alignment = 1 #pour centrer

    # Third Section

    # p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc

#Version de départ
def relance2bis_word(facture,ingeprev,affaire,mission,payeur,envoi):
    self = facture

    doc = Document()
    font = doc.styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(12)
    font = doc.styles['Footer'].font
    font.name = 'Calibri'
    font.size = Pt(8)

    # Entête
    #new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    #new_section.start_type
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    # doc.sections[0].top_margin = Cm(5)

    image = ingeprev.logo
    logo_run = paragraph.add_run()
    # logo_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run.add_picture(image, width=Inches(1.5))
    header.add_paragraph('')
    # doc.add_picture(image, width=Inches(1.5))
    # last_paragraph = doc.paragraphs[-1]
    # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Second Section Entête avec dernières adresses enregistrées dans la base
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[1]

    # Set to 2 column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    new_section = doc.add_section(WD_SECTION.NEW_COLUMN)
    new_section.start_type

    if not adresses_identiques2(payeur, envoi):
        paragraphe = []
        text = '{}'.format(payeur.Denomination_Sociale)
        paragraphe.append(text)
        text = '{}'.format(payeur.Adresse)
        paragraphe.append(text)
        if payeur.Complement_Adresse != '':
            text = '{}'.format(payeur.Complement_Adresse)
            paragraphe.append(text)
        text = '{} {}'.format(payeur.CP, payeur.Ville)
        paragraphe.append(text)
        for text in paragraphe:
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            run = p1.add_run()
            run.add_text(text)

        p1 = doc.add_paragraph()
        p1.paragraph_format.space_before = Pt(20)
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text('s/c')

    paragraphe = []
    text = '{}'.format(envoi.Denomination_Sociale)
    paragraphe.append(text)
    text = '{}'.format(envoi.Adresse)
    paragraphe.append(text)
    if envoi.Complement_Adresse != '':
        text = '{}'.format(envoi.Complement_Adresse)
        paragraphe.append(text)
    text = '{} {}'.format(envoi.CP, envoi.Ville)
    paragraphe.append(text)
    if envoi.Nom_Contact != '':
        text = "A l'attention de {} {} {}".format(envoi.Civilite, envoi.Nom_Contact, envoi.Prenom_Contact)
        paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[3]

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

    paragraphe = []
    text = 'N° de dossier : {}'.format(affaire.Ref_Affaire)
    paragraphe.append(text)
    text = 'Affaire : {}'.format(mission.Nom_Mission)
    paragraphe.append(text)
    text = '{}'.format(mission.Adresse)
    paragraphe.append(text)
    if mission.Complement_Adresse != '':
        text = '{}'.format(mission.Complement_Adresse)
        paragraphe.append(text)
    text = '{} {}'.format(mission.CP, mission.Ville)
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    # p1.alignment = 1 #pour centrer

    # Third Section

    # p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(30)
    p1.paragraph_format.space_after = Pt(30)
    run = p1.add_run()
    run.add_text("Objet : Lettre de relance suite au non-paiement d'une facture")

    paragraphe = ['Madame, Monsieur']
    text = "Après vérification de votre compte client et sauf erreur de notre part, nous n'avons pas perçu le règlement de la facture suivante :"
    paragraphe.append(text)
    text = "Facture n° {} d'un montant de {} € en date du {}, arrivée à échéance le {}".format(self.Numero_Facture,
                                                                                               self.Montant_Facture_TTC(),
                                                                                               self.Date_Facture_aff(),
                                                                                               self.Date_Echeance1_aff())
    paragraphe.append(text)
    if self.Nb_Avoir() == 1:
        avoirlie = self.Avoirs_Lies()[0]
        montantavoir = self.Montants_Avoirs_Lies_TTC()[0]
        text = "En tenant compte de l'avoir partiel n° {} d'un montant de {} € TTC, la somme restant à régler est de {} € TTC.".format(
            avoirlie, montantavoir, self.Reste_A_Payer_TTC())
        paragraphe.append(text)
    if self.Nb_Avoir() > 1:
        avoirlie = self.Avoirs_Lies()
        montantavoir = self.Montants_Avoirs_Lies_TTC()
        text = "En tenant compte des avoirs partiels "
        for x in avoirlie:
            text += 'n°' + str(x) + ', '
        text += " de montants respectifs "
        for x in montantavoir:
            text += str(x) + ' € TTC, '
        text += "la somme restant à régler est de {} € TTC.".format(self.Reste_A_Payer_TTC())
        paragraphe.append(text)
    text = "Pour mémoire, cette facture a déjà dû faire l'objet d'une relance par mail le {}".format(
        self.Date_Relance1_aff())
    paragraphe.append(text)
    text = "Nous vous saurions gré de bien vouloir régulariser cette situation rapidement et de nous confirmer la date du règlement."
    paragraphe.append(text)
    text = "Dans le cas où votre paiement nous parviendrait avant réception de la présente, nous vous prions de ne pas en tenir compte."
    paragraphe.append(text)
    text = "Nous restons à votre disposition pour tout complément d'information."
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(10)
        run = p1.add_run()
        run.add_text(text)

    # Second Section Entête avec dernières adresses enregistrées dans la base
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[4]

    # Set to 2 column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    new_section = doc.add_section(WD_SECTION.NEW_COLUMN)
    new_section.start_type

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(10)
    p1.paragraph_format.space_after = Pt(00)
    run = p1.add_run()
    run.add_text("Pour INGEPREV")
    paragraphe = ["{} {}".format(self.Prenom_Pilote, self.Nom_Pilote)]
    text = "Tél : {}".format(self.Tel_Portable_Pilote)
    paragraphe.append(text)
    text = "Mail : {}".format(self.Email_Pilote)
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.start_type
    section = doc.sections[6]

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(40)
    p1.paragraph_format.space_after = Pt(10)
    run = p1.add_run()
    text = "Pièces jointes : copie de la facture n° : {}".format(self.Numero_Facture)
    run.add_text(text)

    # Calling the footer
    footer = doc.sections[0].footer

    # Calling the paragraph already present in
    # the footer section
    # footer_para = footer.paragraphs[0]
    para = []
    text = 'INGEPREV - {} {} {}\n'.format(ingeprev.Adresse, ingeprev.CP, ingeprev.Ville)
    para.append(text)
    text = '{} au capital de {} € - {} RCS PARIS \n '.format(ingeprev.Type_Societe, ingeprev.Capital,
                                                                 ingeprev.SIRET)
    para.append(text)
    text = '{} - {} - {} - {}'.format(ingeprev.Site_Web, ingeprev.Facebook, ingeprev.Twitter, ingeprev.Linkedin)
    para.append(text)
    for text in para:
        footer_para = footer.add_paragraph()
        run = footer_para.add_run()
        run.font.size = Pt(8)
        footer_para.paragraph_format.space_after = Pt(0)
        footer_para.paragraph_format.space_before = Pt(0)
        run.add_text(text)
    # Adding text in the footer
    # font = footer_para.paragraph_format.font
    # font.size = Pt(8)
    # footer_para.text = text

    return doc


'''
    table = doc.add_table(0,2)
    paragraphe = []
    text = '{}'.format(self.Denomination_Facture)
    paragraphe.append(text)
    text = '{}'.format(self.Adresse_Facture)
    paragraphe.append(text)
    if self.Complement_Adresse_Facture != '':
        text = '{}'.format(self.Complement_Adresse_Facture)
        paragraphe.append(text)
    text = '{} {}'.format(self.CP_Facture, self.Ville_Facture)
    paragraphe.append(text)
    if self.Nom_Facture !='':
        text = "A l'attention de {} {} {}".format(self.Civilite_Facture,self.Nom_Facture,self.Prenom_Facture)
        paragraphe.append(text)
    for text in paragraphe:
        row = table.add_row().cells
        p1 = row[1].paragraphs[0]
        #p1.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    paragraphe = []
    text = 'N° de dossier : {}'.format(affaire.Ref_Affaire)
    paragraphe.append(text)
    text = 'Affaire : {}'.format(mission.Nom_Mission)
    paragraphe.append(text)
    text = '{}'.format(mission.Adresse)
    paragraphe.append(text)
    if mission.Complement_Adresse != '':
        text = '{}'.format(mission.Complement_Adresse)
        paragraphe.append(text)
    text = '{} {}'.format(mission.CP, mission.Ville)
    paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)

    Entête avec adresses liées à la facture
    # Second Section
    #new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    #new_section.start_type
    section = doc.sections[1]

    # Set to 2 column layout
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    new_section = doc.add_section(WD_SECTION.NEW_COLUMN)
    new_section.start_type

    if not adresses_identiques(self):
        paragraphe = []
        text = '{}'.format(self.Denomination_Client)
        paragraphe.append(text)
        text = '{}'.format(self.Adresse_Client)
        paragraphe.append(text)
        if self.Complement_Adresse_Client != '':
            text = '{}'.format(self.Complement_Adresse_Client)
            paragraphe.append(text)
        text = '{} {}'.format(self.CP_Client, self.Ville_Client)
        paragraphe.append(text)
        for text in paragraphe:
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            run = p1.add_run()
            run.add_text(text)

        p1 = doc.add_paragraph()
        p1.paragraph_format.space_before = Pt(20)
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text('s/c')

    paragraphe = []
    text = '{}'.format(self.Denomination_Facture)
    paragraphe.append(text)
    text = '{}'.format(self.Adresse_Facture)
    paragraphe.append(text)
    if self.Complement_Adresse_Facture != '':
        text = '{}'.format(self.Complement_Adresse_Facture)
        paragraphe.append(text)
    text = '{} {}'.format(self.CP_Facture, self.Ville_Facture)
    paragraphe.append(text)
    if self.Nom_Facture != '':
        text = "A l'attention de {} {} {}".format(self.Civilite_Facture, self.Nom_Facture, self.Prenom_Facture)
        paragraphe.append(text)
    for text in paragraphe:
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run()
        run.add_text(text)
    '''
